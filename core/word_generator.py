"""
Word Document Generator for Insurance Policies
Generates Word documents from templates based on Ref_Template_Design configuration
"""
import os
from docx import Document
from docx.shared import Pt
from django.conf import settings
from django.db.models import Model
from core.models import (
    Policy_Main,
    Policy_Main_Product,
    Policy_Main_Product_Item,
    Policy_Main_Product_Item_Risk,
    Policy_Main_Product_Item_Question,
    Ref_Template_Design,
    RefClient,
    Ref_Product,
    Ref_Item,
    Ref_Risk,
    Ref_Item_Question
)
import logging
from decimal import Decimal
from datetime import datetime, date

logger = logging.getLogger(__name__)


def get_field_value(policy, table_name, field_name):
    """
    Get field value from policy-related tables
    
    Args:
        policy: Policy_Main instance
        table_name: Table name (e.g., 'ins_policy_main')
        field_name: Field name (e.g., 'PolicyNo')
    
    Returns:
        Field value or None
    """
    try:
        if table_name == 'ins_policy_main':
            value = getattr(policy, field_name, None)
            # Handle foreign key fields
            if field_name == 'ClientId':
                return policy.ClientId.ClientName if policy.ClientId else None
            elif field_name == 'AgentId':
                return policy.AgentId.username if policy.AgentId else None
            elif field_name == 'PolicyTemplateId':
                return policy.PolicyTemplateId.PolicyTemplateName if policy.PolicyTemplateId else None
            elif field_name == 'CurrencyId':
                return policy.CurrencyId.Currency_name if policy.CurrencyId else None
            elif field_name == 'AgentBranchId':
                return policy.AgentBranchId.BranchName if policy.AgentBranchId else None
            elif field_name == 'AgentChannelId':
                return policy.AgentChannelId.ChannelName if policy.AgentChannelId else None
            elif field_name == 'ApprovedBy':
                return policy.ApprovedBy.username if policy.ApprovedBy else None
            elif field_name == 'CreatedBy':
                return policy.CreatedBy.username if policy.CreatedBy else None
            elif field_name == 'ModifiedBy':
                return policy.ModifiedBy.username if policy.ModifiedBy else None
            return value
        elif table_name == 'ins_policy_main_product':
            # Get first product
            product = Policy_Main_Product.objects.filter(PolicyMainId=policy).select_related('ProductId').first()
            if product:
                if field_name == 'ProductId':
                    return product.ProductId.ProductName if product.ProductId else None
                return getattr(product, field_name, None)
        elif table_name == 'ins_policy_main_product_item':
            # Get first item
            product = Policy_Main_Product.objects.filter(PolicyMainId=policy).first()
            if product:
                item = Policy_Main_Product_Item.objects.filter(PolicyMainProductId=product).select_related('ItemId').first()
                if item:
                    if field_name == 'ItemId':
                        return item.ItemId.ItemName if item.ItemId else None
                    return getattr(item, field_name, None)
        elif table_name == 'ins_policy_main_product_item_risk':
            # Get first risk
            product = Policy_Main_Product.objects.filter(PolicyMainId=policy).first()
            if product:
                item = Policy_Main_Product_Item.objects.filter(PolicyMainProductId=product).first()
                if item:
                    risk = Policy_Main_Product_Item_Risk.objects.filter(PolicyMainProductItemId=item).select_related('RiskId').first()
                    if risk:
                        if field_name == 'RiskId':
                            return risk.RiskId.RiskName if risk.RiskId else None
                        return getattr(risk, field_name, None)
        elif table_name == 'ins_policy_main_product_item_question':
            # Get first question
            product = Policy_Main_Product.objects.filter(PolicyMainId=policy).first()
            if product:
                item = Policy_Main_Product_Item.objects.filter(PolicyMainProductId=product).first()
                if item:
                    question = Policy_Main_Product_Item_Question.objects.filter(PolicyMainProductItemId=item).select_related('ItemQuestionId').first()
                    if question:
                        if field_name == 'ItemQuestionId':
                            return question.ItemQuestionId.ItemQuestionName if question.ItemQuestionId else None
                        return getattr(question, field_name, None)
    except Exception as e:
        logger.error(f'Error getting field value for {table_name}.{field_name}: {str(e)}')
    return None


def get_field_values_list(policy, table_name, field_name):
    """
    Get list of field values for dynamic fields (IsStatic=True)
    
    Args:
        policy: Policy_Main instance
        table_name: Table name
        field_name: Field name
    
    Returns:
        List of field values
    """
    values = []
    try:
        if table_name == 'ins_policy_main_product':
            products = Policy_Main_Product.objects.filter(PolicyMainId=policy).select_related('ProductId')
            for product in products:
                if field_name == 'ProductId':
                    value = product.ProductId.ProductName if product.ProductId else None
                else:
                    value = getattr(product, field_name, None)
                if value is not None:
                    values.append(value)
        elif table_name == 'ins_policy_main_product_item':
            product = Policy_Main_Product.objects.filter(PolicyMainId=policy).first()
            if product:
                items = Policy_Main_Product_Item.objects.filter(PolicyMainProductId=product).select_related('ItemId')
                for item in items:
                    if field_name == 'ItemId':
                        value = item.ItemId.ItemName if item.ItemId else None
                    else:
                        value = getattr(item, field_name, None)
                    if value is not None:
                        values.append(value)
        elif table_name == 'ins_policy_main_product_item_risk':
            product = Policy_Main_Product.objects.filter(PolicyMainId=policy).first()
            if product:
                item = Policy_Main_Product_Item.objects.filter(PolicyMainProductId=product).first()
                if item:
                    risks = Policy_Main_Product_Item_Risk.objects.filter(PolicyMainProductItemId=item).select_related('RiskId')
                    for risk in risks:
                        if field_name == 'RiskId':
                            value = risk.RiskId.RiskName if risk.RiskId else None
                        else:
                            value = getattr(risk, field_name, None)
                        if value is not None:
                            values.append(value)
        elif table_name == 'ins_policy_main_product_item_question':
            product = Policy_Main_Product.objects.filter(PolicyMainId=policy).first()
            if product:
                item = Policy_Main_Product_Item.objects.filter(PolicyMainProductId=product).first()
                if item:
                    questions = Policy_Main_Product_Item_Question.objects.filter(PolicyMainProductItemId=item).select_related('ItemQuestionId')
                    for question in questions:
                        if field_name == 'ItemQuestionId':
                            value = question.ItemQuestionId.ItemQuestionName if question.ItemQuestionId else None
                        else:
                            value = getattr(question, field_name, None)
                        if value is not None:
                            values.append(value)
    except Exception as e:
        logger.error(f'Error getting field values list for {table_name}.{field_name}: {str(e)}')
    return values


def format_field_value(value):
    """
    Format field value for display in Word document
    
    Args:
        value: Field value (can be various types)
    
    Returns:
        Formatted string
    """
    if value is None:
        return ''
    if isinstance(value, bool):
        return 'Тийм' if value else 'Үгүй'
    if isinstance(value, (date, datetime)):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, Decimal):
        return str(value)
    return str(value)


def replace_placeholder_in_paragraph(paragraph, placeholder, value):
    """
    Replace placeholder in a paragraph, handling multiple occurrences
    
    Args:
        paragraph: docx paragraph object
        placeholder: Placeholder text (e.g., '[PolicyNo]')
        value: Replacement value
    """
    if placeholder in paragraph.text:
        # Get original text and formatting before clearing
        full_text = paragraph.text
        original_runs = paragraph.runs
        original_font_name = original_runs[0].font.name if original_runs else None
        original_font_size = original_runs[0].font.size if original_runs else None
        
        # Clear paragraph
        paragraph.clear()
        
        # Split by placeholder and rebuild
        parts = full_text.split(placeholder)
        for i, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                if original_font_name:
                    run.font.name = original_font_name
                if original_font_size:
                    run.font.size = original_font_size
            if i < len(parts) - 1:
                # Add the replacement value
                run = paragraph.add_run(str(value))
                if original_font_name:
                    run.font.name = original_font_name
                if original_font_size:
                    run.font.size = original_font_size


def replace_placeholder_in_table(table, placeholder, value):
    """
    Replace placeholder in all cells of a table
    Optimized to prevent freezing
    
    Args:
        table: docx table object
        placeholder: Placeholder text
        value: Replacement value
    """
    value_str = str(value)
    max_cells = 10000  # Safety limit
    cell_count = 0
    
    for row in table.rows:
        for cell in row.cells:
            if cell_count >= max_cells:
                logger.warning(f'Reached maximum cell limit ({max_cells}) for table replacement')
                return
                
            for paragraph in cell.paragraphs:
                if placeholder in paragraph.text:
                    try:
                        # Get original formatting
                        original_runs = paragraph.runs
                        original_font_name = original_runs[0].font.name if original_runs and len(original_runs) > 0 else None
                        original_font_size = original_runs[0].font.size if original_runs and len(original_runs) > 0 else None
                        
                        # Clear and rebuild
                        full_text = paragraph.text
                        paragraph.clear()
                        
                        parts = full_text.split(placeholder)
                        for i, part in enumerate(parts):
                            if part:
                                run = paragraph.add_run(part)
                                if original_font_name:
                                    run.font.name = original_font_name
                                if original_font_size:
                                    run.font.size = original_font_size
                            if i < len(parts) - 1:
                                run = paragraph.add_run(value_str)
                                if original_font_name:
                                    run.font.name = original_font_name
                                if original_font_size:
                                    run.font.size = original_font_size
                    except Exception as e:
                        logger.warning(f'Error replacing in table cell paragraph: {str(e)}')
                        # Fallback to simple replacement
                        try:
                            paragraph.text = paragraph.text.replace(placeholder, value_str)
                        except:
                            pass  # Skip if still fails
            cell_count += 1


def _placeholder_exists_in_document(doc, placeholder):
    """
    Check if placeholder exists in document (paragraphs or tables)
    
    Args:
        doc: Document object
        placeholder: Placeholder text to search for
    
    Returns:
        True if placeholder exists, False otherwise
    """
    # Check paragraphs
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            return True
    
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        return True
    return False


def replace_text_in_document(doc, placeholder, value):
    """
    Replace placeholder text throughout the entire document
    Optimized to prevent freezing on large documents
    
    Args:
        doc: Document object
        placeholder: Placeholder text (e.g., '[PolicyNo]')
        value: Replacement value
    """
    value_str = str(value)
    
    # Replace in all paragraphs (limit to prevent freezing)
    paragraph_count = 0
    max_paragraphs = 10000  # Safety limit
    
    for paragraph in doc.paragraphs:
        if paragraph_count >= max_paragraphs:
            logger.warning(f'Reached maximum paragraph limit ({max_paragraphs}) for replacement')
            break
            
        if placeholder in paragraph.text:
            try:
                # Preserve formatting
                runs = paragraph.runs
                if runs and len(runs) > 0:
                    # Get formatting from first run
                    font_name = runs[0].font.name if runs[0].font.name else None
                    font_size = runs[0].font.size if runs[0].font.size else None
                    
                    # Clear and rebuild
                    full_text = paragraph.text
                    paragraph.clear()
                    
                    parts = full_text.split(placeholder)
                    for i, part in enumerate(parts):
                        if part:
                            run = paragraph.add_run(part)
                            if font_name:
                                run.font.name = font_name
                            if font_size:
                                run.font.size = font_size
                        if i < len(parts) - 1:
                            run = paragraph.add_run(value_str)
                            if font_name:
                                run.font.name = font_name
                            if font_size:
                                run.font.size = font_size
                else:
                    # Simple replacement if no runs
                    paragraph.text = paragraph.text.replace(placeholder, value_str)
            except Exception as e:
                logger.warning(f'Error replacing in paragraph: {str(e)}')
                # Fallback to simple replacement
                try:
                    paragraph.text = paragraph.text.replace(placeholder, value_str)
                except:
                    pass  # Skip if still fails
        paragraph_count += 1
    
    # Replace in all tables (limit to prevent freezing)
    table_count = 0
    max_tables = 1000  # Safety limit
    
    for table in doc.tables:
        if table_count >= max_tables:
            logger.warning(f'Reached maximum table limit ({max_tables}) for replacement')
            break
        try:
            replace_placeholder_in_table(table, placeholder, value_str)
        except Exception as e:
            logger.warning(f'Error replacing in table: {str(e)}')
        table_count += 1


def generate_policy_word_document(policy_id):
    """
    Generate Word document for a policy based on template design
    
    Args:
        policy_id: Policy ID
    
    Returns:
        Path to generated Word document or None if error
    """
    try:
        # Get policy with related objects
        policy = Policy_Main.objects.select_related(
            'PolicyTemplateId', 'ClientId', 'AgentId', 'CurrencyId',
            'AgentBranchId', 'AgentChannelId', 'ApprovedBy', 'CreatedBy', 'ModifiedBy'
        ).prefetch_related(
            'policy_products__ProductId',
            'policy_products__product_items__ItemId',
            'policy_products__product_items__item_risks__RiskId',
            'policy_products__product_items__item_questions__ItemQuestionId'
        ).get(PolicyId=policy_id)
        
        # Get template
        template = policy.PolicyTemplateId
        
        # Check if template has Word file
        if not template.FilePath:
            error_msg = f'Word template file path not set for template "{template.PolicyTemplateName}" (ID: {template.PolicyTemplateId}). Please set the file path in template management.'
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        # Check if file exists (handle both absolute and relative paths)
        template_path = template.FilePath
        original_path = template_path
        
        if not os.path.isabs(template_path):
            # Relative path - check in media root or project root
            from pathlib import Path
            base_dir = Path(settings.BASE_DIR)
            media_path = os.path.join(settings.MEDIA_ROOT, template_path)
            base_path = os.path.join(base_dir, template_path)
            
            # Try different locations
            possible_paths = [
                media_path,
                base_path,
                template_path,  # Try as-is
                os.path.join(base_dir, 'templates', 'word_templates', os.path.basename(template_path)),
                os.path.join(settings.MEDIA_ROOT, 'templates', 'word_templates', os.path.basename(template_path))
            ]
            
            found_path = None
            for path in possible_paths:
                if os.path.exists(path):
                    found_path = path
                    break
            
            if found_path:
                template_path = found_path
            else:
                error_msg = f'Word template file not found. Searched in:\n'
                error_msg += f'- Original path: {original_path}\n'
                error_msg += f'- Media path: {media_path}\n'
                error_msg += f'- Base path: {base_path}\n'
                error_msg += f'- Template: {template.PolicyTemplateName} (ID: {template.PolicyTemplateId})'
                logger.error(error_msg)
                raise FileNotFoundError(f'Template file not found: {original_path}')
        
        if not os.path.exists(template_path):
            error_msg = f'Word template file not found: {template_path} (Original: {original_path})'
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
        
        # Load Word template
        doc = Document(template_path)
        
        # Get all template designs for this template
        template_designs = Ref_Template_Design.objects.filter(
            PolicyTemplateId=template,
            IsActive=True
        ).order_by('TableNameEng', 'FieldNameEng')
        
        # Process each template design
        # Primary format: [TableName.FieldName]
        # Limit processing to prevent freezing on large documents
        max_replacements = 1000  # Safety limit
        replacement_count = 0
        
        for design in template_designs:
            if replacement_count >= max_replacements:
                logger.warning(f'Reached maximum replacement limit ({max_replacements}) for policy {policy_id}')
                break
                
            # Use format [TableName.FieldName] for placeholders
            placeholder = f'[{design.TableNameEng}.{design.FieldNameEng}]'
            
            if design.IsStatic:
                # Single value - replace throughout document
                value = get_field_value(policy, design.TableNameEng, design.FieldNameEng)
                formatted_value = format_field_value(value)
                
                # Replace in entire document
                replace_text_in_document(doc, placeholder, formatted_value)
                replacement_count += 1
            else:
                # Multiple rows - get all values
                values = get_field_values_list(policy, design.TableNameEng, design.FieldNameEng)
                
                if values:
                    # Format all values
                    formatted_values = [format_field_value(v) for v in values]
                    
                    # For dynamic fields, join with newline or comma
                    # Limit the number of values to prevent freezing
                    if len(formatted_values) > 100:
                        formatted_values = formatted_values[:100]
                        logger.warning(f'Truncated dynamic field values to 100 for {placeholder}')
                    
                    combined_value = '\n'.join(formatted_values)
                    
                    # Replace in entire document
                    replace_text_in_document(doc, placeholder, combined_value)
                    replacement_count += 1
                else:
                    # No values found - replace with empty string
                    replace_text_in_document(doc, placeholder, '')
                    replacement_count += 1
        
        # Also support legacy format [FieldNameEng] for backward compatibility
        # Only process if legacy placeholders might exist
        for design in template_designs:
            if replacement_count >= max_replacements:
                break
                
            legacy_placeholder = f'[{design.FieldNameEng}]'
            
            # Check if legacy placeholder exists in document
            if _placeholder_exists_in_document(doc, legacy_placeholder):
                if design.IsStatic:
                    value = get_field_value(policy, design.TableNameEng, design.FieldNameEng)
                    formatted_value = format_field_value(value)
                    replace_text_in_document(doc, legacy_placeholder, formatted_value)
                    replacement_count += 1
                else:
                    values = get_field_values_list(policy, design.TableNameEng, design.FieldNameEng)
                    if values:
                        formatted_values = [format_field_value(v) for v in values]
                        if len(formatted_values) > 100:
                            formatted_values = formatted_values[:100]
                        combined_value = '\n'.join(formatted_values)
                        replace_text_in_document(doc, legacy_placeholder, combined_value)
                        replacement_count += 1
                    else:
                        replace_text_in_document(doc, legacy_placeholder, '')
                        replacement_count += 1
        
        # Save generated document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'policy_documents')
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate unique filename
        output_filename = f'policy_{policy.PolicyNo}_{policy.PolicyId}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        output_path = os.path.join(output_dir, output_filename)
        
        doc.save(output_path)
        
        logger.info(f'Word document generated successfully: {output_path}')
        return output_path
        
    except Policy_Main.DoesNotExist:
        logger.error(f'Policy {policy_id} not found')
        return None
    except Exception as e:
        logger.error(f'Error generating Word document for policy {policy_id}: {str(e)}', exc_info=True)
        return None

